#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档转换工具
功能：将包含图片的Markdown文档转换为PDF和Word格式
作者：Assistant
路径：/document_converter.py
"""

import os
import re
import sys
import logging
from pathlib import Path
from typing import List, Dict, Tuple
import argparse

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('document_converter.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def install_requirements():
    """【依赖管理】安装必要的依赖包"""
    import subprocess
    
    required_packages = [
        'python-docx',
        'markdown',
        'weasyprint',  # 用于PDF生成
        'Pillow',      # 图片处理
        'beautifulsoup4',  # HTML解析
        'lxml'         # XML解析
    ]
    
    logger.info("检查并安装必要的依赖包...")
    for package in required_packages:
        try:
            __import__(package.replace('-', '_'))
            logger.info(f"✓ {package} 已安装")
        except ImportError:
            logger.info(f"安装 {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

class MarkdownToDocConverter:
    """【模块化设计】Markdown到文档转换器"""
    
    def __init__(self, markdown_file: str, output_dir: str = "output"):
        """
        初始化转换器
        
        Args:
            markdown_file: Markdown文件路径
            output_dir: 输出目录
        """
        self.markdown_file = Path(markdown_file)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # 确保图片路径存在
        self.base_dir = self.markdown_file.parent
        
        logger.info(f"初始化转换器：{self.markdown_file} -> {self.output_dir}")
    
    def read_markdown(self) -> str:
        """【文件操作】读取Markdown文件内容"""
        try:
            with open(self.markdown_file, 'r', encoding='utf-8') as f:
                content = f.read()
            logger.info(f"成功读取Markdown文件，长度：{len(content)} 字符")
            return content
        except Exception as e:
            logger.error(f"读取Markdown文件失败：{e}")
            raise
    
    def process_images(self, content: str) -> Tuple[str, List[str]]:
        """【图片处理】处理图片路径，返回处理后的内容和图片路径列表"""
        # 匹配Markdown图片语法：![alt](path)
        image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        images = []
        project_root = Path.cwd()
        agent_root = project_root / 'agent'
        base_parts = list(self.base_dir.parts)
        base_files_prefix = None
        if 'files' in base_parts:
            try:
                files_idx = base_parts.index('files')
                base_files_prefix = '/'.join(base_parts[files_idx:])
            except ValueError:
                base_files_prefix = None

        def resolve_image_path(image_path_str: str) -> Path:
            """在多候选根目录下解析图片路径，返回第一个存在的绝对路径"""
            norm = image_path_str.replace('\\', '/')
            p = Path(norm)
            if p.is_absolute() and p.exists():
                return p

            candidates = []
            # 1) 相对当前markdown文件所在目录
            candidates.append(self.base_dir / norm)
            # 2) 相对项目根目录
            candidates.append(project_root / norm)
            # 3) 相对agent根目录
            candidates.append(agent_root / norm)
            # 4) 若路径以 base_files_prefix/ 开头（如 files/sam/company_xxx/...），裁剪该前缀后拼到 base_dir
            if base_files_prefix and norm.startswith(base_files_prefix + '/'):
                suffix = norm[len(base_files_prefix) + 1:]
                candidates.append(self.base_dir / suffix)
            # 5) 若以 files/ 开头，优先考虑 agent 根目录
            if norm.startswith('files/'):
                candidates.append(agent_root / norm)

            for c in candidates:
                if c.exists():
                    return c.resolve()

            # 兜底：返回 base_dir 下的绝对化路径（即使不存在，用于日志展示）
            return (self.base_dir / norm).resolve()

        def replace_image(match):
            alt_text = match.group(1)
            image_path = match.group(2)
            full_path = resolve_image_path(image_path)
            if full_path and full_path.exists():
                images.append(str(full_path))
                logger.info(f"找到图片：{full_path}")
                return f'![{alt_text}]({full_path})'
            else:
                logger.warning(f"图片不存在（已尝试多路径解析）：{image_path} -> {full_path}")
                return f'[图片缺失: {alt_text}]'

        processed_content = re.sub(image_pattern, replace_image, content)
        return processed_content, images
    
    def convert_to_word(self, content: str) -> str:
        """【文档生成】转换为Word文档"""
        try:
            from docx import Document
            from docx.shared import Inches, RGBColor, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            import markdown
            from bs4 import BeautifulSoup
            
            logger.info("开始转换为Word文档...")
            
            # 创建Word文档
            doc = Document()
            
            # 【设计规范】设置默认字体为宋体
            def set_font_style(element, font_name="宋体", font_size=12, color=None, bold=False):
                """设置字体样式"""
                for paragraph in element if hasattr(element, '__iter__') else [element]:
                    if hasattr(paragraph, 'runs'):
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                            if color:
                                run.font.color.rgb = color
                            run.font.bold = bold
                            # 设置中文字体
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            
            # 处理图片路径
            processed_content, images = self.process_images(content)
            
            # 将Markdown转换为HTML
            html = markdown.markdown(processed_content, extensions=['tables', 'codehilite'])
            soup = BeautifulSoup(html, 'html.parser')
            
            # 解析HTML并添加到Word文档
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'img']):
                if element.name.startswith('h'):
                    # 【设计规范】标题设置 - 黑色宋体
                    level = int(element.name[1])
                    heading = doc.add_heading(element.get_text(), level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 设置标题字体样式
                    font_sizes = {1: 18, 2: 16, 3: 14, 4: 13, 5: 12, 6: 12}
                    font_size = font_sizes.get(level, 12)
                    
                    for run in heading.runs:
                        run.font.name = "宋体"
                        run.font.size = Pt(font_size)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                        run.font.bold = True
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                
                elif element.name == 'p':
                    # 【设计规范】段落设置 - 宋体
                    para = doc.add_paragraph(element.get_text())
                    set_font_style(para, font_name="宋体", font_size=12, color=RGBColor(0, 0, 0))
                
                elif element.name in ['ul', 'ol']:
                    # 【设计规范】列表设置 - 宋体
                    for li in element.find_all('li'):
                        list_para = doc.add_paragraph(li.get_text(), style='List Bullet')
                        set_font_style(list_para, font_name="宋体", font_size=12, color=RGBColor(0, 0, 0))
                
                elif element.name == 'img':
                    # 【设计规范】图片处理
                    try:
                        src = element.get('src')
                        alt = element.get('alt', '图片')
                        
                        if src and os.path.exists(src):
                            # 添加图片标题
                            caption_para = doc.add_paragraph(f"图片：{alt}", style='Caption')
                            set_font_style(caption_para, font_name="宋体", font_size=10, color=RGBColor(0, 0, 0))
                            # 添加图片
                            doc.add_picture(src, width=Inches(6))
                            logger.info(f"添加图片到Word：{src}")
                        else:
                            missing_para = doc.add_paragraph(f"[图片缺失: {alt}]")
                            set_font_style(missing_para, font_name="宋体", font_size=12, color=RGBColor(255, 0, 0))
                    except Exception as e:
                        logger.error(f"添加图片到Word时出错：{e}")
                        error_para = doc.add_paragraph(f"[图片处理错误: {element.get('alt', '未知图片')}]")
                        set_font_style(error_para, font_name="宋体", font_size=12, color=RGBColor(255, 0, 0))
            
            # 【设计规范】设置文档默认样式
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0)
            
            # 保存Word文档
            word_output = self.output_dir / f"{self.markdown_file.stem}.docx"
            doc.save(word_output)
            logger.info(f"Word文档保存成功：{word_output}")
            return str(word_output)
            
        except Exception as e:
            logger.error(f"转换为Word文档失败：{e}")
            raise
    
    def convert_to_pdf_html(self, content: str) -> str:
        """【文档生成】通过HTML中间格式转换为PDF"""
        try:
            import markdown
            from weasyprint import HTML, CSS
            
            logger.info("开始转换为PDF文档（通过HTML）...")
            
            # 处理图片路径
            processed_content, images = self.process_images(content)
            
            # 将Markdown转换为HTML
            html_content = markdown.markdown(
                processed_content, 
                extensions=['tables', 'codehilite', 'toc']
            )
            
            # 创建完整的HTML文档
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>{self.markdown_file.stem}</title>
                <style>
                    body {{
                        font-family: "Microsoft YaHei", "SimSun", Arial, sans-serif;
                        line-height: 1.6;
                        margin: 40px;
                        color: #333;
                    }}
                    h1, h2, h3, h4, h5, h6 {{
                        color: #2c3e50;
                        margin-top: 2em;
                        margin-bottom: 1em;
                    }}
                    h1 {{ font-size: 2.2em; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                    h2 {{ font-size: 1.8em; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; }}
                    h3 {{ font-size: 1.4em; }}
                    p {{ margin-bottom: 1em; text-align: justify; }}
                    img {{ 
                        max-width: 100%; 
                        height: auto; 
                        display: block; 
                        margin: 20px auto;
                        border: 1px solid #ddd;
                        border-radius: 4px;
                        padding: 5px;
                    }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin: 20px 0;
                    }}
                    th, td {{
                        border: 1px solid #ddd;
                        padding: 12px;
                        text-align: left;
                    }}
                    th {{
                        background-color: #f2f2f2;
                        font-weight: bold;
                    }}
                    code {{
                        background-color: #f8f9fa;
                        padding: 2px 4px;
                        border-radius: 3px;
                        font-family: "Consolas", "Monaco", monospace;
                    }}
                    pre {{
                        background-color: #f8f9fa;
                        padding: 15px;
                        border-radius: 5px;
                        overflow-x: auto;
                    }}
                    blockquote {{
                        border-left: 4px solid #3498db;
                        margin: 20px 0;
                        padding-left: 20px;
                        color: #7f8c8d;
                    }}
                    ul, ol {{
                        margin-bottom: 1em;
                        padding-left: 30px;
                    }}
                    li {{
                        margin-bottom: 0.5em;
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # 保存HTML文件（用于调试）
            html_output = self.output_dir / f"{self.markdown_file.stem}.html"
            with open(html_output, 'w', encoding='utf-8') as f:
                f.write(full_html)
            logger.info(f"HTML文件保存：{html_output}")
            
            # 转换为PDF
            pdf_output = self.output_dir / f"{self.markdown_file.stem}.pdf"
            HTML(string=full_html, base_url=str(self.base_dir)).write_pdf(str(pdf_output))
            
            logger.info(f"PDF文档保存成功：{pdf_output}")
            return str(pdf_output)
            
        except Exception as e:
            logger.error(f"转换为PDF文档失败：{e}")
            raise
    
    def convert_all(self) -> Dict[str, str]:
        """【业务逻辑】转换为所有支持的格式"""
        results = {}
        content = self.read_markdown()
        
        try:
            # 转换为Word
            word_file = self.convert_to_word(content)
            results['word'] = word_file
            logger.info(f"✓ Word转换完成：{word_file}")
        except Exception as e:
            logger.error(f"✗ Word转换失败：{e}")
            results['word'] = None
        
        try:
            # 转换为PDF
            pdf_file = self.convert_to_pdf_html(content)
            results['pdf'] = pdf_file
            logger.info(f"✓ PDF转换完成：{pdf_file}")
        except Exception as e:
            logger.error(f"✗ PDF转换失败：{e}")
            results['pdf'] = None
        
        return results

def main():
    """【测试策略】主函数，处理命令行参数"""
    parser = argparse.ArgumentParser(description='将Markdown文档转换为PDF和Word格式')
    parser.add_argument('input_file', help='输入的Markdown文件路径')
    parser.add_argument('-o', '--output', default='output', help='输出目录（默认：output）')
    parser.add_argument('--install-deps', action='store_true', help='安装必要的依赖包')
    
    args = parser.parse_args()
    
    if args.install_deps:
        install_requirements()
        return
    
    # 检查输入文件
    if not os.path.exists(args.input_file):
        logger.error(f"输入文件不存在：{args.input_file}")
        return
    
    try:
        # 创建转换器并执行转换
        converter = MarkdownToDocConverter(args.input_file, args.output)
        results = converter.convert_all()
        
        # 输出结果
        print("\n" + "="*60)
        print("文档转换完成！")
        print("="*60)
        
        for format_type, file_path in results.items():
            if file_path:
                print(f"✓ {format_type.upper():>6}: {file_path}")
            else:
                print(f"✗ {format_type.upper():>6}: 转换失败")
        
        print("="*60)
        
    except Exception as e:
        logger.error(f"转换过程中发生错误：{e}")
        sys.exit(1)

if __name__ == "__main__":
    main() 