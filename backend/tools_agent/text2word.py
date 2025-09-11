#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件：text2word.py
路径：/tools_agent/text2word.py
功能：将Markdown文本转换为Word文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import re
import os
import sys
import io
import tempfile

# 增加一个简单的日志记录器
def log_message(level, message):
    print(f"[{level}] [text2word] {message}", file=sys.stderr)

def create_safe_document():
    """
    安全地创建Document对象，绕过所有模板依赖
    """
    try:
        # 方法1: 直接创建
        log_message("INFO", "尝试直接创建Document...")
        doc = Document()
        log_message("INFO", "直接创建Document成功")
        return doc
    except Exception as e1:
        log_message("WARNING", f"直接创建Document失败: {e1}")
        
        try:
            # 方法2: 使用None参数
            log_message("INFO", "尝试使用Document(None)...")
            doc = Document(None)
            log_message("INFO", "Document(None)创建成功")
            return doc
        except Exception as e2:
            log_message("WARNING", f"Document(None)失败: {e2}")
            
            try:
                # 方法3: 创建临时空白文档
                log_message("INFO", "尝试创建临时空白文档...")
                
                # 创建一个最小的空白docx文件内容
                minimal_docx_content = create_minimal_docx_bytes()
                
                # 写入临时文件
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                    tmp_file.write(minimal_docx_content)
                    tmp_path = tmp_file.name
                
                # 使用临时文件创建Document
                doc = Document(tmp_path)
                
                # 清理临时文件
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                
                log_message("INFO", "临时空白文档创建成功")
                return doc
                
            except Exception as e3:
                log_message("ERROR", f"所有Document创建方法都失败了: {e3}")
                raise e3

def create_minimal_docx_bytes():
    """
    创建一个最小的有效docx文件的字节内容
    """
    import zipfile
    import io
    
    # 创建内存中的zip文件
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # 添加必需的文件
        
        # 1. [Content_Types].xml
        content_types = '''<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''
        zip_file.writestr('[Content_Types].xml', content_types)
        
        # 2. _rels/.rels
        rels = '''<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''
        zip_file.writestr('_rels/.rels', rels)
        
        # 3. word/document.xml
        document_xml = '''<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:sectPr>
<w:pgSz w:w="12240" w:h="15840"/>
<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>
</w:sectPr>
</w:body>
</w:document>'''
        zip_file.writestr('word/document.xml', document_xml)
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def text_to_word(markdown_text, doc_path):
    """
    将Markdown文本转换为Word文档
    
    Args:
        markdown_text (str): 要转换的Markdown文本
        doc_path (str): 输出的Word文档路径
    """
    try:
        log_message("INFO", "开始创建Word文档...")
        
        # 使用安全的文档创建方法
        doc = create_safe_document()
        log_message("INFO", "成功创建Word文档")

        # 设置文档基本样式
        try:
            if hasattr(doc, 'styles') and 'Normal' in doc.styles:
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                style.font.size = Pt(12)
                log_message("INFO", "成功设置文档基本样式")
        except Exception as e:
                log_message("WARNING", f"设置文档样式失败，使用默认样式: {e}")

    except Exception as e:
        log_message("ERROR", f"创建Document对象时出错: {e}")
        raise e

    def add_paragraph_custom(text, style=None, font_name='宋体', size=12, paragraph_format=None, line_spacing=Pt(24), is_bold=False):
        p = doc.add_paragraph(style=style)
        run = p.add_run(text)
        run.bold = is_bold
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        if paragraph_format:
            p.paragraph_format.alignment = paragraph_format
        p.paragraph_format.line_spacing = line_spacing
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.space_before = Pt(8)

    def add_code_block(code_text, font_name='Courier New', size=10):
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(100, 100, 100)
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(6)
        
    def process_text_with_citations(text, paragraph):
        # 匹配引用模式，如[1]或[1, 2, 3]
        current_pos = 0
        while True:
            # 查找下一个引用
            match = re.search(r'\[(\d+(?:,\s*\d+)*)\]', text[current_pos:])
            if not match:
                # 没有更多引用，添加剩余文本
                if current_pos < len(text):
                    run = paragraph.add_run(text[current_pos:])
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(12)
                break
            
            # 添加引用前的文本
            if match.start() > 0:
                run = paragraph.add_run(text[current_pos:current_pos + match.start()])
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
            
            # 添加引用
            citation_text = match.group(1)
            run = paragraph.add_run(f'[{citation_text}]')
            run.font.superscript = True
            run.font.size = Pt(10)
            
            # 更新位置
            current_pos += match.end()
            
    def process_cell_text(cell, text):
        # 清空单元格中的默认段落
        for p in cell.paragraphs:
            p.clear()
        
        # 创建新段落
        p = cell.paragraphs[0]
        
        # 处理单元格中的粗体文本和引用
        if '**' in text:
            parts = text.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # 粗体部分
                    run = p.add_run(part)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(12)
                else:  # 非粗体部分
                    process_text_with_citations(part, p)
        else:
            process_text_with_citations(text, p)

    def parse_table_rows(table_lines):
        """解析Markdown表格行并返回结构化数据"""
        rows_data = []
        for line in table_lines:
            if line.strip() and line.strip().startswith('|') and line.strip().endswith('|'):
                # 分割并清理单元格数据
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                rows_data.append(cells)
        
        if len(rows_data) < 3:  # 至少需要表头、分隔行和一行数据
            return None
            
        # 移除分隔行（通常是第二行，包含 ----- 的行）
        header = rows_data[0]
        data_rows = rows_data[2:]
        
        return header, data_rows

    def process_table(table_lines):
        # 解析表格数据
        parsed_data = parse_table_rows(table_lines)
        if not parsed_data:
            return None
        
        header, data_rows = parsed_data
        
        # 创建表格
        table = doc.add_table(rows=len(data_rows) + 1, cols=len(header))
        
        # 尝试设置表格样式（如果可用）
        try:
            table.style = 'Table Grid'
        except Exception as e:
            log_message("WARNING", f"无法设置表格样式: {e}")
        
        # 设置表格宽度
        try:
            table.autofit = False
            table_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
            table.width = table_width
        except Exception as e:
            log_message("WARNING", f"无法设置表格宽度: {e}")
        
        # 填充表头
        for j, header_text in enumerate(header):
            cell = table.cell(0, j)
            try:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                log_message("WARNING", f"无法设置表头对齐: {e}")
            process_cell_text(cell, header_text)
            # 设置表头加粗
            for run in cell.paragraphs[0].runs:
                run.bold = True
        
        # 填充表格内容
        for i, row_data in enumerate(data_rows):
            for j, cell_text in enumerate(row_data):
                if j < len(header):  # 确保不超出列数
                    cell = table.cell(i + 1, j)
                    process_cell_text(cell, cell_text)
        
        # 分析表格分隔行中的对齐信息
        alignment_row = table_lines[1]
        alignments = []
        if alignment_row.strip().startswith('|') and alignment_row.strip().endswith('|'):
            alignment_cells = [cell.strip() for cell in alignment_row.split('|')[1:-1]]
            for align_text in alignment_cells:
                if align_text.startswith(':') and align_text.endswith(':'):
                    alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                elif align_text.startswith(':'):
                    alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
                elif align_text.endswith(':'):
                    alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
                else:
                    alignments.append(WD_ALIGN_PARAGRAPH.LEFT)  # 默认左对齐
        
        # 应用对齐方式
        if alignments:
            try:
                for i in range(len(data_rows) + 1):  # 包括表头行
                    for j, alignment in enumerate(alignments):
                        if j < len(header):
                            for paragraph in table.cell(i, j).paragraphs:
                                paragraph.alignment = alignment
            except Exception as e:
                log_message("WARNING", f"无法设置表格对齐: {e}")
        
        # 添加表格后的空行
        doc.add_paragraph()
        return table

    log_message("INFO", "开始处理Markdown文本...")
    
    lines = markdown_text.split("\n")
    in_code_block = False
    in_table = False
    code_block = []
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 处理代码块
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                add_code_block("\n".join(code_block), size=10)
                code_block = []
            i += 1
            continue
        
        if in_code_block:
            code_block.append(line)
            i += 1
            continue
        
        # 处理表格
        if line.strip().startswith("|") and not in_table:
            # 检查下一行是否是分隔行（包含 ----- 的行）
            if i + 1 < len(lines) and re.match(r'\s*\|[\s\-:\|]+\|\s*$', lines[i + 1]):
                in_table = True
                table_lines = [line]
                i += 1
                continue
        
        if in_table:
            table_lines.append(line)
            # 检查是否是表格的最后一行
            if i + 1 >= len(lines) or not lines[i + 1].strip().startswith("|"):
                # 表格结束
                process_table(table_lines)
                table_lines = []
                in_table = False
            i += 1
            continue
        
        # 处理标题和普通段落
        try:
            if line.startswith("# "):
                add_paragraph_custom(line[2:].strip(), 'Heading 1', size=18, paragraph_format=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=Pt(28))
            elif line.startswith("## "):
                add_paragraph_custom(line[3:].strip(), 'Heading 2', size=16, paragraph_format=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=Pt(26))
            elif line.startswith("### "):
                add_paragraph_custom(line[4:].strip(), 'Heading 3', size=14, paragraph_format=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=Pt(24))
            elif line.startswith("#### "):
                add_paragraph_custom(line[5:].strip(), 'Heading 4', size=14, paragraph_format=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=Pt(24))
            elif line.strip() != "":
                # 处理包含引用的段落
                parts = line.split('**')
                p = doc.add_paragraph()
                
                if len(parts) > 1:  # 有粗体标记
                    for j, part in enumerate(parts):
                        if j % 2 == 1:  # 粗体部分
                            run = p.add_run(part)
                            run.bold = True
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.size = Pt(12)
                        else:  # 非粗体部分
                            process_text_with_citations(part, p)
                else:  # 没有粗体标记，直接处理引用
                    process_text_with_citations(line, p)
                
                p.paragraph_format.line_spacing = Pt(24)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.space_before = Pt(8)
        except Exception as e:
            log_message("ERROR", f"处理文本行时出错: {line[:50]}... 错误: {e}")
            # 如果样式处理失败，至少添加纯文本
            try:
                p = doc.add_paragraph(line)
            except Exception as e2:
                log_message("ERROR", f"添加纯文本段落也失败: {e2}")
        
        i += 1

    try:
        log_message("INFO", f"准备保存文档到: {doc_path}")
        # 确保目标目录存在
        target_dir = os.path.dirname(doc_path)
        if target_dir and not os.path.exists(target_dir):
            os.makedirs(target_dir)
            log_message("INFO", f"创建目标目录: {target_dir}")
        
        doc.save(doc_path)
        log_message("INFO", f"成功保存DOCX文件: {doc_path}")
    except Exception as e:
        log_message("ERROR", f"保存DOCX文件失败 {doc_path}: {e}")
        # 向上抛出异常，让调用者知道保存失败了
        raise


# 测试代码
if __name__ == "__main__":
    # 示例文本
    example_markdown = """
# 中英文混合文档示例

This is an example of a mixed Chinese and English document. 这是一个中英文混合文档的示例。

## 格式测试

1. **Bold text** and 常规文本
2. 中文**粗体**测试
3. 引用测试[1]和多引用[2, 3, 4]

### 表格测试

下面是一个示例表格：

| 表头1   | 表头2      | 表头3      |
| :------ | :--------: | ---------: |
| 内容1   | 内容2      | 内容3      |
| **粗体**| 引用内容[5]| 普通内容   |

### 代码块示例

以下是一个Python代码块：

```python
def hello_world():
    print("Hello, 世界！")
```

### 图片示例

下面是一个示例图片：

![示例图片](https://example.com/image.jpg)
"""

    # 调用函数
    text_to_word(example_markdown, "files/test.docx")

